#!/usr/bin/env python3
"""
Generate codebook documents for RDS dataset files.

This script reads RDS files from the specified directory and generates
Word documents (.docx) following the codebook template format.

Usage:
    python generate_codebook.py [--input-dir PATH] [--output-dir PATH] [--single FILE]

Requirements:
    - python-docx (mamba install -n thesis python-docx)
    - R (mamba install -n thesis r-base)
"""

import argparse
import csv
import json
import os
import subprocess
import sys
from collections import defaultdict
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT


# Custom descriptions for derived variables not in V_METADATA.csv
# These are computed during data extraction and analysis
CUSTOM_DESCRIPTIONS = {
    # Identifiers and dates
    'bdate': 'Birth date (mid-year approximation)',
    'diagn_date': 'Date of first eligible depression diagnosis',
    'prescr': 'Date of first antidepressant prescription after diagnosis',
    'date_death': 'Date of death',
    'date_emig': 'Date of emigration',
    'date_fail': 'Date of suicidal behavior outcome event',
    'admin_end': 'Administrative end of follow-up (2020-12-31)',

    # Demographics
    'age': 'Age at diagnosis date (years)',
    'age_diagn': 'Age at diagnosis date (years)',
    'agecat': 'Age category at diagnosis (6-17 or 18-24)',
    'female': 'Sex indicator (0=male, 1=female)',

    # Treatment variables
    'cc': 'Treatment group (1=initiator within 28 days, 0=non-initiator)',
    'predi_diff': 'Days between diagnosis and first prescription',
    'switch': 'Treatment switch indicator (1=switched to SSRI after 28 days)',
    'switch_dat': 'Date of treatment switch',
    'ssri': 'SSRI treatment indicator',
    'exp': 'Antidepressant exposure indicator (time-varying)',

    # Follow-up variables
    'fu_start': 'Start of follow-up period',
    'fu_end12': 'End of 12-week follow-up period',
    'fu_end52': 'End of 52-week follow-up period',
    'fu_end_pp': 'End of per-protocol follow-up (censored at switch/death/emigration)',
    'fu_end_itt': 'End of intention-to-treat follow-up (censored at death/emigration)',
    'fu_lenght': 'Length of ITT follow-up in days',

    # Outcome variables
    'sb12': 'Suicidal behavior within 12 weeks (1=yes, 0=no)',
    'sb52': 'Suicidal behavior within 52 weeks (1=yes, 0=no)',
    'sb12_pp': 'Suicidal behavior within 12 weeks, per-protocol (1=yes, 0=no)',
    'sb12_itt': 'Suicidal behavior within 12 weeks, intention-to-treat (1=yes, 0=no)',
    'sb52_pp': 'Suicidal behavior within 52 weeks, per-protocol (1=yes, 0=no)',
    'sb52_itt': 'Suicidal behavior within 52 weeks, intention-to-treat (1=yes, 0=no)',

    # Censoring variables
    'cens_death': 'Censored due to death (1=yes, 0=no)',
    'cens_emig': 'Censored due to emigration (1=yes, 0=no)',
    'cens_switch': 'Censored due to treatment switch (1=yes, 0=no)',
    'cens_admin': 'Censored at administrative end date (1=yes, 0=no)',
    'cens_deathemig': 'Censored due to death or emigration (1=yes, 0=no)',

    # Family history covariates
    'fh_suicidal': 'Family history of suicidal behavior (0=no, 1=yes, 2=unknown)',
    'fh_depr': 'Family history of depression (0=no, 1=yes, 2=unknown)',

    # Socioeconomic covariates
    'edufam_cat': 'Highest parental education level (0=primary, 1=secondary, 2=tertiary, 99=unknown)',
    'inc_cat': 'Family income category (1=low, 2=middle, 3=high, NOINFO=unknown)',
    'inc': 'Family disposable income',
    '_p20': '20th percentile of income distribution for year',
    '_p80': '80th percentile of income distribution for year',

    # Psychiatric diagnosis covariates (before follow-up)
    'bipolar': 'Prior bipolar disorder diagnosis (F30-F31)',
    'anxiety': 'Prior anxiety disorder diagnosis (F4x)',
    'psychotic': 'Prior psychotic disorder diagnosis (F2x)',
    'sud': 'Prior substance use disorder diagnosis (F11-F19)',
    'eat': 'Prior eating disorder diagnosis (F50)',
    'personality': 'Prior personality disorder diagnosis (F60-F61)',
    'adhd': 'Prior ADHD diagnosis (F90)',
    'developmental': 'Prior developmental disorder diagnosis (F70-F98, excl F84, F90, F91)',
    'asd': 'Prior autism spectrum disorder diagnosis (F84)',
    'cd': 'Prior conduct disorder diagnosis (F91)',
    'alc': 'Prior alcohol use disorder diagnosis (F10)',
    'pois_gen': 'Prior poisoning diagnosis (X40-X49, T36-T50, excl alcohol)',
    'pois_alc': 'Prior alcohol poisoning diagnosis (X45, T510)',
    'suicide': 'Prior suicide attempt/self-harm diagnosis (X60-X84, Y10-Y34)',
    'chapterV': 'Any prior ICD-10 Chapter V mental disorder diagnosis (F00-F99)',

    # Medication covariates (90 days before follow-up)
    'apsych': 'Prior antipsychotic medication (N05A, excl lithium)',
    'hypno': 'Prior hypnotic/sedative medication (N05B/N05C, excl benzodiazepines)',
    'benzo': 'Prior benzodiazepine anxiolytic medication (N05BA)',
    'anepi': 'Prior antiepileptic medication (N03A)',
    'stimul': 'Prior psychostimulant medication (N06B)',
    'drugs': 'Prior drugs for addictive disorders (N07B)',
    'opioid': 'Prior opioid medication (N02A)',

    # Time-varying medication indicators
    'apsych_tv': 'Antipsychotic use in current period (time-varying)',
    'benzo_tv': 'Benzodiazepine use in current period (time-varying)',
    'opsych_tv': 'Other psychotropic use in current period (time-varying)',
    'opsych': 'Other psychotropic medication indicator (not antidepressant, antipsychotic, or benzodiazepine)',
    'anypsych_tv': 'Any psychotropic use in current period (apsych_tv OR benzo_tv OR opsych_tv)',

    # Hospitalization variables
    'hosp': 'Number of prior psychiatric hospitalizations (0, 1, 2-3, 4+)',
    'hosp_no': 'Count of prior psychiatric hospitalizations',
    'stay_time': 'Length of hospital stay in days',

    # Time-varying cohort variables
    'week': 'Week number within follow-up period',
    'period_start': 'Start date of current period (weekly)',
    'period_end': 'End date of current period (weekly)',
    'start': 'Start date of medication treatment episode',
    'end': 'End date of medication treatment episode',

    # Outcome event variables
    'intent': 'Intent of suicidal behavior (known or unknown)',
    'keep': 'Record retention indicator',
    'date': 'Event/prescription date',

    # Diagnosis source variables
    'year': 'Calendar year of diagnosis',

    # Prescription-related
    'otherprescr': 'Date of other (non-SSRI) psychotropic prescription',
    'min_diff': 'Minimum days between prescriptions',
    'max_diff': 'Maximum days between prescriptions',
}


def load_variable_descriptions(metadata_csv: Path) -> dict:
    """Load variable descriptions from V_METADATA.csv and custom descriptions.

    Returns a dict mapping lowercase variable names to their descriptions.
    Custom descriptions take priority over V_METADATA.csv descriptions.
    When multiple descriptions exist for a variable in V_METADATA.csv, the most common one is used.
    """
    # Start with custom descriptions
    result = {k.lower(): v for k, v in CUSTOM_DESCRIPTIONS.items()}

    # Load from V_METADATA.csv
    descriptions = defaultdict(list)

    with open(metadata_csv, 'r', encoding='latin-1') as f:
        reader = csv.DictReader(f)
        for row in reader:
            col = row.get('COLUMN_NAME', '').strip()
            desc = row.get('COL_DESCRIPTION_ENG', '').strip()
            if col and desc:
                descriptions[col.lower()].append(desc)

    # For each variable, pick the most common description (if not already in custom)
    for var, descs in descriptions.items():
        if var not in result:  # Custom descriptions take priority
            # Count occurrences and pick most common
            desc_counts = defaultdict(int)
            for d in descs:
                desc_counts[d] += 1
            most_common = max(desc_counts.keys(), key=lambda x: desc_counts[x])
            result[var] = most_common

    return result


def extract_rds_metadata(rds_path: Path) -> dict:
    """Extract metadata from an RDS file using R."""
    r_script = f'''
    df <- readRDS("{rds_path}")
    if (!is.data.frame(df)) {{
        stop("RDS file does not contain a data frame")
    }}

    col_info <- data.frame(
        variable = names(df),
        class = sapply(df, function(x) paste(class(x), collapse="/")),
        n_unique = sapply(df, function(x) length(unique(x))),
        n_na = sapply(df, function(x) sum(is.na(x))),
        n_rows = nrow(df),
        stringsAsFactors = FALSE
    )

    col_info$sample_values <- sapply(names(df), function(col) {{
        x <- df[[col]]
        uniq <- unique(x[!is.na(x)])
        if (length(uniq) <= 10) {{
            paste(head(sort(as.character(uniq)), 10), collapse=", ")
        }} else {{
            paste(head(as.character(uniq), 5), collapse=", ")
        }}
    }})

    library(jsonlite)
    cat(toJSON(col_info, pretty = FALSE))
    '''

    result = subprocess.run(
        ['Rscript', '-e', r_script],
        capture_output=True,
        text=True
    )

    if result.returncode != 0:
        raise RuntimeError(f"R script failed: {result.stderr}")

    metadata = json.loads(result.stdout)
    return metadata


def r_class_to_type_format(r_class: str, n_unique: int, sample_values: str) -> str:
    """Convert R class to human-readable type/format."""
    class_lower = r_class.lower()

    if 'date' in class_lower:
        return 'Date (YYYY-MM-DD)'
    elif 'posixct' in class_lower or 'posixlt' in class_lower:
        return 'DateTime'
    elif class_lower in ('integer', 'numeric'):
        if n_unique <= 10:
            return 'Numeric (categorical)'
        return 'Numeric'
    elif class_lower == 'character':
        if n_unique <= 20:
            return 'Character (categorical)'
        return 'Character'
    elif class_lower == 'factor':
        return 'Factor'
    elif class_lower == 'logical':
        return 'Logical (TRUE/FALSE)'
    else:
        return r_class


def generate_value_description(r_class: str, n_unique: int, n_na: int, n_rows: int, sample_values: str) -> str:
    """Generate value description based on metadata."""
    parts = []

    # Show sample values for categorical variables
    if n_unique <= 10 and sample_values:
        values = sample_values.split(', ')
        if r_class.lower() in ('integer', 'numeric'):
            parts.append(f"Values: {sample_values}")
        else:
            parts.append(f"Values: {sample_values}")
    elif sample_values:
        parts.append(f"Examples: {sample_values}...")

    # Add uniqueness info
    if n_unique == n_rows:
        parts.append("(unique identifier)")
    elif n_unique <= 10:
        parts.append(f"({n_unique} distinct values)")
    else:
        parts.append(f"({n_unique:,} distinct values)")

    # Add missing data info
    if n_na > 0:
        pct_missing = (n_na / n_rows) * 100
        parts.append(f"Missing: {n_na:,} ({pct_missing:.1f}%)")

    return '\n'.join(parts)


def create_codebook_document(dataset_name: str, metadata: list, output_path: Path,
                             var_descriptions: dict = None):
    """Create a codebook document for a dataset.

    Args:
        dataset_name: Name of the dataset
        metadata: List of variable metadata dicts
        output_path: Path to save the document
        var_descriptions: Optional dict mapping variable names to descriptions
    """
    if var_descriptions is None:
        var_descriptions = {}

    doc = Document()

    # Title
    title = doc.add_heading(f'Codebook: {dataset_name}', level=0)

    # Description paragraph
    n_rows = metadata[0]['n_rows'] if metadata else 0
    n_cols = len(metadata)
    desc = doc.add_paragraph()
    desc.add_run(f'This codebook provides explanations and formats for all variables in the data set ')
    desc.add_run(dataset_name).bold = True
    desc.add_run(f'.')

    # Dataset summary
    summary = doc.add_paragraph()
    summary.add_run(f'Dataset contains {n_rows:,} observations and {n_cols} variables.')

    doc.add_paragraph()  # Spacer

    # Create table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_cells = table.rows[0].cells
    headers = ['VARIABLE NAME', 'DESCRIPTION', 'TYPE/FORMAT', 'VALUE DESCRIPTION', 'USED IN ANALYSES', 'COMMENTS']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for var_info in metadata:
        row_cells = table.add_row().cells

        variable = var_info['variable']
        r_class = var_info['class']
        n_unique = var_info['n_unique']
        n_na = var_info['n_na']
        n_rows = var_info['n_rows']
        sample_values = var_info.get('sample_values', '')

        row_cells[0].text = variable
        row_cells[1].text = var_descriptions.get(variable.lower(), '')
        row_cells[2].text = r_class_to_type_format(r_class, n_unique, sample_values)
        row_cells[3].text = generate_value_description(r_class, n_unique, n_na, n_rows, sample_values)
        row_cells[4].text = ''  # Used in analyses - to be filled manually
        row_cells[5].text = ''  # Comments - to be filled manually

        # Format cells
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.2)  # Variable name
        row.cells[1].width = Inches(1.5)  # Description
        row.cells[2].width = Inches(1.0)  # Type/Format
        row.cells[3].width = Inches(2.0)  # Value Description
        row.cells[4].width = Inches(0.8)  # Used in analyses
        row.cells[5].width = Inches(1.0)  # Comments

    # Save document
    doc.save(output_path)
    print(f"Created: {output_path}")


def main():
    parser = argparse.ArgumentParser(description='Generate codebook documents for RDS files')
    parser.add_argument('--input-dir', type=Path,
                        default=Path('suicidality/extraction/output/rds'),
                        help='Directory containing RDS files')
    parser.add_argument('--output-dir', type=Path,
                        default=Path('suicidality/Documents/codebook'),
                        help='Output directory for codebook documents')
    parser.add_argument('--single', type=str,
                        help='Generate codebook for a single RDS file only')
    parser.add_argument('--use-cached-metadata', action='store_true',
                        help='Use cached metadata JSON instead of reading RDS files')
    parser.add_argument('--metadata-json', type=Path,
                        default=Path('suicidality/Documents/rds_metadata.json'),
                        help='Path to cached metadata JSON file')
    parser.add_argument('--var-descriptions', type=Path,
                        default=Path('data/V_METADATA.csv'),
                        help='Path to variable descriptions CSV file')

    args = parser.parse_args()

    # Load variable descriptions
    var_descriptions = {}
    if args.var_descriptions.exists():
        print(f"Loading variable descriptions from {args.var_descriptions}")
        var_descriptions = load_variable_descriptions(args.var_descriptions)
        print(f"Loaded descriptions for {len(var_descriptions)} variables")
    else:
        print(f"Warning: Variable descriptions file not found: {args.var_descriptions}")

    # Ensure output directory exists
    args.output_dir.mkdir(parents=True, exist_ok=True)

    if args.use_cached_metadata and args.metadata_json.exists():
        print(f"Using cached metadata from {args.metadata_json}")
        with open(args.metadata_json) as f:
            all_metadata = json.load(f)

        for filename, metadata in all_metadata.items():
            if args.single and filename != args.single:
                continue

            dataset_name = filename.replace('.rds', '')
            output_path = args.output_dir / f'{dataset_name}.docx'
            create_codebook_document(dataset_name, metadata, output_path, var_descriptions)
    else:
        # Find all RDS files
        rds_files = list(args.input_dir.glob('*.rds'))

        # Skip files that are not data frames
        skip_files = {'cohort_flow_summary.rds'}
        rds_files = [f for f in rds_files if f.name not in skip_files]

        if not rds_files:
            print(f"No RDS files found in {args.input_dir}")
            sys.exit(1)

        print(f"Found {len(rds_files)} RDS files")

        for rds_path in sorted(rds_files):
            if args.single and rds_path.name != args.single:
                continue

            print(f"\nProcessing: {rds_path.name}")
            try:
                metadata = extract_rds_metadata(rds_path)
                dataset_name = rds_path.stem
                output_path = args.output_dir / f'codebook_{dataset_name}.docx'
                create_codebook_document(dataset_name, metadata, output_path, var_descriptions)
            except Exception as e:
                print(f"  Error: {e}")
                continue

    print(f"\nCodebooks generated in: {args.output_dir}")


if __name__ == '__main__':
    main()

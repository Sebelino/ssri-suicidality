#!/usr/bin/env python3
"""
Generate a logbook document documenting the data extraction pipeline.

This script creates a Word document (.docx) following the MEB logbook template format,
documenting the data flow from source databases to final analysis datasets.

Usage:
    python generate_logbook.py [--output PATH]

Requirements:
    - python-docx (mamba install -n thesis python-docx)
    - graphviz (mamba install -n thesis graphviz)
"""

import argparse
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Data extraction pipeline documentation
# Scripts 01-11: Database extraction (require database access)
# Scripts 12-24: RDS processing (no database access)
# Format: (script_name, inputs, outputs, description)
EXTRACTION_PIPELINE = [
    # Database extraction scripts (01-11)
    (
        "01_raw_diagnoses_index.R",
        ["v_npr_dia (database)"],
        ["raw_diagnoses_index.rds"],
        "Extract depression diagnoses (F32-F33) from 2006-2019, excluding remission codes (F334)."
    ),
    (
        "02_raw_prescriptions_all.R",
        ["v_lmr (database)"],
        ["raw_prescriptions.rds"],
        "Extract all psychotropic prescriptions (N06A, N02A, N03A, N05A, N05B, N05C, N06B, N07B) for washout and covariate extraction."
    ),
    (
        "03_raw_individual_bootstrap.R",
        ["v_individual (database)", "raw_diagnoses_index.rds"],
        ["raw_individual.rds"],
        "Extract birth dates and sex for individuals with depression diagnoses."
    ),
    (
        "04_define_cohort.R",
        ["raw_diagnoses_index.rds", "raw_prescriptions.rds", "raw_individual.rds", "v_parent (database)"],
        ["cohort_lopnrs.rds", "parent_lopnrs.rds", "cohort_base.rds", "cohort_flow_summary.rds"],
        "Define cohort: filter by age 6-24, apply 365-day antidepressant washout (N06A), identify SSRI prescriptions (N06AB), extract parent links. Includes both initiators and non-initiators."
    ),
    (
        "05_raw_migration.R",
        ["v_migration (database)", "cohort_lopnrs.rds"],
        ["raw_migration.rds"],
        "Extract emigration dates for cohort members."
    ),
    (
        "06_raw_dor.R",
        ["v_dor_bas (database)", "v_dor_orsak (database)", "cohort_lopnrs.rds"],
        ["raw_dor.rds"],
        "Extract death dates and causes of death for cohort members."
    ),
    (
        "07_raw_diagnoses_cohort.R",
        ["v_npr_dia (database)", "cohort_lopnrs.rds"],
        ["raw_diagnoses_cohort.rds"],
        "Extract all relevant diagnoses for cohort: psychiatric (F*), self-harm (X60-X84), undetermined intent (Y10-Y34), poisoning (T36-T50, X40-X49)."
    ),
    (
        "08_raw_diagnoses_parents.R",
        ["v_npr_dia (database)", "parent_lopnrs.rds"],
        ["raw_diagnoses_parents.rds"],
        "Extract depression and suicidal behavior diagnoses for parents (family history covariates)."
    ),
    (
        "09_raw_lisa.R",
        ["v_lisa (database)", "parent_lopnrs.rds"],
        ["raw_lisa.rds"],
        "Extract LISA data (education, income) for parents."
    ),
    (
        "10_raw_hospitalization.R",
        ["v_npr_dia (database)", "cohort_lopnrs.rds"],
        ["raw_hospitalization.rds"],
        "Extract inpatient hospitalization data for censoring (psychiatric stays >= 2 days)."
    ),
    (
        "11_raw_prescriptions_cohort.R",
        ["v_lmr (database)", "cohort_lopnrs.rds"],
        ["raw_prescriptions_cohort.rds"],
        "Extract all prescriptions for cohort members (for time-varying medication analysis)."
    ),
    # RDS processing scripts (12-24)
    (
        "12_process_base.R",
        ["cohort_base.rds", "raw_migration.rds", "raw_dor.rds"],
        ["base_28.rds"],
        "Create base cohort with death and emigration dates added."
    ),
    (
        "13_process_outcomes.R",
        ["raw_diagnoses_cohort.rds", "raw_dor.rds"],
        ["dia_all_28.rds"],
        "Identify suicidal behavior outcomes: suicide attempts (X60-X84, Y10-Y34) and deaths from suicide."
    ),
    (
        "14_process_censoring.R",
        ["raw_hospitalization.rds"],
        ["cens_hosp_28.rds"],
        "Identify hospitalization censoring events (psychiatric stays >= 2 days)."
    ),
    (
        "15_process_followup.R",
        ["base_28.rds", "dia_all_28.rds", "cens_hosp_28.rds"],
        ["main_12wks_28_tmp.rds"],
        "Create main 12-week follow-up cohort. Assign treatment groups (initiators vs non-initiators). Calculate outcomes and censoring."
    ),
    (
        "16_process_cov_family_history.R",
        ["cohort_base.rds", "raw_diagnoses_parents.rds"],
        ["cov_family_history.rds"],
        "Create family history covariates: parental depression and suicidal behavior."
    ),
    (
        "17_process_cov_education.R",
        ["cohort_base.rds", "raw_lisa.rds"],
        ["cov_education.rds"],
        "Create education covariate: highest parental education level."
    ),
    (
        "18_process_cov_income.R",
        ["cohort_base.rds", "raw_lisa.rds"],
        ["cov_income.rds"],
        "Create income covariate: family income categorized by percentiles."
    ),
    (
        "19_process_cov_diagnoses.R",
        ["main_12wks_28_tmp.rds", "raw_diagnoses_cohort.rds"],
        ["cov_diagnoses.rds"],
        "Create diagnosis covariates: prior psychiatric diagnoses (bipolar, anxiety, psychotic, SUD, etc.)."
    ),
    (
        "20_process_cov_medications.R",
        ["main_12wks_28_tmp.rds", "raw_prescriptions_cohort.rds"],
        ["cov_medications.rds", "othermeds_28.rds"],
        "Create medication covariates: prior psychotropic medications within 90 days."
    ),
    (
        "21_process_cov_hospitalizations.R",
        ["main_12wks_28_tmp.rds", "raw_hospitalization.rds"],
        ["cov_hospitalizations.rds"],
        "Create hospitalization covariate: count of prior psychiatric hospitalizations."
    ),
    (
        "22_process_covariates_assembly.R",
        ["base_28.rds", "cov_family_history.rds", "cov_education.rds", "cov_income.rds", "cov_diagnoses.rds", "cov_medications.rds", "cov_hospitalizations.rds"],
        ["base_cov_28.rds"],
        "Assemble all covariates into final covariate dataset."
    ),
    (
        "23_process_time_varying.R",
        ["main_12wks_28_tmp.rds", "othermeds_28.rds", "raw_prescriptions_cohort.rds"],
        ["pp_12wks_max_tmp.rds"],
        "Create time-varying medication episodes and per-protocol cohort with weekly intervals. Censors initiators who discontinue treatment."
    ),
    (
        "24_process_final_cohorts.R",
        ["main_12wks_28_tmp.rds", "pp_12wks_max_tmp.rds", "base_cov_28.rds", "raw_individual.rds", "raw_diagnoses_index.rds"],
        ["main_12wks_28.rds", "pp_12wks_max.rds"],
        "Create final analysis datasets: join cohorts with covariates, add sex, year, source coding."
    ),
]

# Final output datasets (highlighted in flow chart)
FINAL_DATASETS = {
    "main_12wks_28.rds",
    "pp_12wks_max.rds",
}

# Scripts and outputs to exclude from DAG (none currently)
EXCLUDED_FROM_DAG = set()

# Database tables used
DATABASE_TABLES = {
    "v_lmr": "Swedish Prescribed Drug Register (Läkemedelsregistret)",
    "v_npr_dia": "National Patient Register - Diagnoses",
    "v_individual": "Individual register (demographics)",
    "v_migration": "Migration register (emigration dates)",
    "v_dor_bas": "Cause of Death Register - Basic data",
    "v_dor_orsak": "Cause of Death Register - Causes",
    "v_parent": "Multi-Generation Register (parent links)",
    "v_lisa": "LISA register (education, income)",
}


def generate_flowchart(output_path: Path) -> bool:
    """Generate a flow chart of the data pipeline using Graphviz.

    Args:
        output_path: Path to save the PNG image

    Returns:
        True if successful, False otherwise
    """
    # Filter out excluded scripts from the pipeline
    pipeline = [
        (script, inputs, outputs, desc)
        for script, inputs, outputs, desc in EXTRACTION_PIPELINE
        if script not in EXCLUDED_FROM_DAG
    ]
    # Also filter out excluded outputs from remaining scripts
    pipeline = [
        (script, inputs, [o for o in outputs if o not in EXCLUDED_FROM_DAG], desc)
        for script, inputs, outputs, desc in pipeline
    ]

    # Collect all unique databases
    databases = set()
    for _, inputs, _, _ in pipeline:
        for inp in inputs:
            if "(database)" in inp:
                db_name = inp.replace(" (database)", "")
                databases.add(db_name)

    # Collect all RDS files and build producer map (which script produces which RDS)
    all_rds = set()
    rds_producer = {}  # rds_file -> script that produces it
    for script, inputs, outputs, _ in pipeline:
        for inp in inputs:
            if inp.endswith('.rds'):
                all_rds.add(inp)
        for out in outputs:
            if out.endswith('.rds'):
                all_rds.add(out)
                rds_producer[out] = script

    # Compute topological depth for each node
    # Databases are at depth 0
    # Scripts are at depth = max(depth of inputs) + 1
    # RDS files are at depth = depth of producing script + 1
    node_depth = {}

    # Initialize databases at depth 0
    for db in databases:
        node_depth[db] = 0

    # Process scripts in pipeline order (they're already topologically sorted)
    for script, inputs, outputs, _ in pipeline:
        max_input_depth = 0
        for inp in inputs:
            inp_name = inp.replace(" (database)", "")
            if inp_name in node_depth:
                max_input_depth = max(max_input_depth, node_depth[inp_name])
        # Script is one level below its deepest input
        node_depth[script] = max_input_depth + 1
        # Outputs are one level below the script
        for out in outputs:
            node_depth[out] = node_depth[script] + 1

    # Group nodes by depth
    depth_to_nodes = {}
    for node, depth in node_depth.items():
        if depth not in depth_to_nodes:
            depth_to_nodes[depth] = []
        depth_to_nodes[depth].append(node)

    # Build the DOT graph
    dot_lines = [
        'digraph DataPipeline {',
        '    rankdir=TB;',
        '    newrank=true;',
        '    node [shape=box, style=filled, fontname="Helvetica", fontsize=10];',
        '    edge [fontname="Helvetica", fontsize=8];',
        '',
        '    // Database nodes (cylinder shape)',
        '    node [shape=cylinder, fillcolor="#E8E8E8"];',
    ]

    for db in sorted(databases):
        dot_lines.append(f'    "{db}" [label="{db}"];')

    dot_lines.append('')
    dot_lines.append('    // RDS file nodes')
    dot_lines.append('    node [shape=box, fillcolor="#FFFACD"];')

    # Mark final datasets with different color
    for rds in sorted(all_rds):
        if rds in FINAL_DATASETS:
            dot_lines.append(f'    "{rds}" [fillcolor="#90EE90", style="filled,bold"];')
        else:
            dot_lines.append(f'    "{rds}";')

    dot_lines.append('')
    dot_lines.append('    // Script nodes (rounded)')
    dot_lines.append('    node [shape=box, style="filled,rounded", fillcolor="#ADD8E6"];')

    for script, _, _, _ in pipeline:
        short_name = script.replace(".R", "")
        dot_lines.append(f'    "{script}" [label="{short_name}"];')

    dot_lines.append('')
    dot_lines.append('    // Edges')

    # Create edges
    for script, inputs, outputs, _ in pipeline:
        for inp in inputs:
            inp_name = inp.replace(" (database)", "")
            dot_lines.append(f'    "{inp_name}" -> "{script}";')
        for out in outputs:
            dot_lines.append(f'    "{script}" -> "{out}";')

    # Add rank constraints for layered layout based on computed depths
    dot_lines.append('')
    dot_lines.append('    // Rank constraints by topological depth')
    for depth in sorted(depth_to_nodes.keys()):
        nodes = depth_to_nodes[depth]
        if nodes:
            node_list = '; '.join(f'"{n}"' for n in sorted(nodes))
            dot_lines.append(f'    {{ rank=same; {node_list}; }}')

    dot_lines.append('')
    dot_lines.append('    // Legend')
    dot_lines.append('    subgraph cluster_legend {')
    dot_lines.append('        label="Legend";')
    dot_lines.append('        fontsize=10;')
    dot_lines.append('        style=dashed;')
    dot_lines.append('        node [shape=box, width=1.2, height=0.3];')
    dot_lines.append('        leg_db [shape=cylinder, label="Database", fillcolor="#E8E8E8"];')
    dot_lines.append('        leg_rds [label="RDS file", fillcolor="#FFFACD"];')
    dot_lines.append('        leg_final [label="Final dataset", fillcolor="#90EE90", style="filled,bold"];')
    dot_lines.append('        leg_script [label="R script", style="filled,rounded", fillcolor="#ADD8E6"];')
    dot_lines.append('        leg_db -> leg_rds -> leg_final -> leg_script [style=invis];')
    dot_lines.append('    }')

    dot_lines.append('}')

    dot_content = '\n'.join(dot_lines)

    # Write DOT file and generate PNG
    try:
        with tempfile.NamedTemporaryFile(mode='w', suffix='.dot', delete=False) as f:
            f.write(dot_content)
            dot_path = f.name

        # Run graphviz
        result = subprocess.run(
            ['dot', '-Tpng', '-Gdpi=150', dot_path, '-o', str(output_path)],
            capture_output=True,
            text=True
        )

        if result.returncode != 0:
            print(f"Graphviz error: {result.stderr}")
            return False

        # Clean up temp file
        Path(dot_path).unlink()

        return True

    except FileNotFoundError:
        print("Warning: Graphviz 'dot' command not found. Install with: mamba install -n thesis graphviz")
        return False
    except Exception as e:
        print(f"Error generating flow chart: {e}")
        return False


def create_logbook_document(output_path: Path, project_name: str = "SSRI-Suicidality"):
    """Create a logbook document documenting the data extraction pipeline."""
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run(f"LOGBOOK_{project_name}.docx")
    title_run.bold = True
    title_run.font.size = Pt(14)

    # Metadata
    today = datetime.now().strftime("%Y-%m-%d")
    doc.add_paragraph(f"Created: {today}")
    doc.add_paragraph(f"Updated: {today}")
    doc.add_paragraph()

    # Description
    desc = doc.add_paragraph()
    desc.add_run("This file describes how the datasets in the ").italic = False
    desc.add_run(project_name).bold = True
    desc.add_run(" project have been created. The project investigates the effect of SSRI treatment on suicidal behavior risk in young people with depression.")

    doc.add_paragraph()

    # Section: LOGBOOK
    logbook_title = doc.add_paragraph()
    logbook_title.add_run("LOGBOOK").bold = True
    logbook_title.paragraph_format.space_after = Pt(12)

    # File folders section
    folders = doc.add_paragraph()
    folders.add_run("File folders:").bold = True

    doc.add_paragraph("Unless otherwise specified:")
    doc.add_paragraph("  - R extraction scripts are in: suicidality/extraction/")
    doc.add_paragraph("  - Output RDS files are in: suicidality/extraction/output/rds/")
    doc.add_paragraph("  - Analysis scripts are in: suicidality/analysis/")
    doc.add_paragraph()

    # Database tables section
    db_title = doc.add_paragraph()
    db_title.add_run("Database tables used:").bold = True

    for table_name, description in DATABASE_TABLES.items():
        doc.add_paragraph(f"  - {table_name}: {description}")

    doc.add_paragraph()

    # Data extraction pipeline section
    pipeline_title = doc.add_paragraph()
    pipeline_title.add_run("Data extraction pipeline:").bold = True
    pipeline_title.paragraph_format.space_after = Pt(12)

    # Pipeline description
    doc.add_paragraph("Scripts are numbered by execution order (01-24):")
    doc.add_paragraph("  - Scripts 01-11: Database extraction (require database access)")
    doc.add_paragraph("  - Scripts 12-24: RDS processing (no database access)")

    doc.add_paragraph()

    # Create main table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    headers = ['Input Data', 'R Script', 'Output Data', 'Description']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for script, inputs, outputs, description in EXTRACTION_PIPELINE:
        row_cells = table.add_row().cells

        row_cells[0].text = '\n'.join(inputs)
        row_cells[1].text = script
        row_cells[2].text = '\n'.join(outputs)
        row_cells[3].text = description

        # Format cells
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(1.5)
        row.cells[2].width = Inches(1.5)
        row.cells[3].width = Inches(2.7)

    doc.add_paragraph()

    # Final datasets section
    final_title = doc.add_paragraph()
    final_title.add_run("Final analysis datasets:").bold = True

    final_datasets = [
        ("main_12wks_28.rds", "Main ITT cohort with 12-week follow-up, includes all covariates"),
        ("pp_12wks_max.rds", "Per-protocol cohort with 12-week follow-up, weekly time-varying treatment"),
    ]

    for dataset, description in final_datasets:
        para = doc.add_paragraph()
        para.add_run(f"  - {dataset}: ").bold = True
        para.add_run(description)

    doc.add_paragraph()

    # Key variables section
    vars_title = doc.add_paragraph()
    vars_title.add_run("Key variables:").bold = True

    key_vars = [
        ("cc", "Treatment group (1=SSRI initiator within 28 days of diagnosis, 0=non-initiator). Uses N06AB (SSRIs only) per Lagerberg 2023."),
        ("sb12", "Suicidal behavior outcome within 12 weeks"),
        ("sb12_pp", "Suicidal behavior outcome (per-protocol analysis)"),
        ("sb12_itt", "Suicidal behavior outcome (intention-to-treat analysis)"),
        ("fu_start", "Start of follow-up period"),
        ("fu_end_pp", "End of per-protocol follow-up (censored at treatment discontinuation)"),
        ("fu_end_itt", "End of intention-to-treat follow-up"),
        ("exp", "Time-varying SSRI exposure indicator (per-protocol cohort)"),
    ]

    for var, description in key_vars:
        para = doc.add_paragraph()
        para.add_run(f"  - {var}: ").bold = True
        para.add_run(description)

    doc.add_paragraph()

    # Methodology notes
    method_title = doc.add_paragraph()
    method_title.add_run("Methodology notes:").bold = True

    notes = [
        "Treatment assignment uses 28-day grace period with SSRI prescriptions only (N06AB) per Lagerberg et al. 2023",
        "Cohort includes both initiators and non-initiators (target trial emulation)",
        "Washout period (365 days) uses all antidepressants (N06A) to exclude prior users",
        "Suicidal behavior defined as ICD-10 codes X60-X84 (known intent) and Y10-Y34 (unknown intent)",
        "Outcomes include both hospital visits for suicide attempts and deaths from suicide",
        "Per-protocol analysis censors initiators at treatment discontinuation",
        "Time-varying medication exposure uses prescription-based treatment periods",
    ]

    for note in notes:
        doc.add_paragraph(f"  - {note}")

    doc.add_paragraph()

    # Flow chart section
    flowchart_title = doc.add_paragraph()
    flowchart_title.add_run("Data extraction flow chart:").bold = True
    flowchart_title.paragraph_format.space_after = Pt(12)

    # Generate flow chart image
    flowchart_path = output_path.parent / "pipeline_flowchart.png"
    if generate_flowchart(flowchart_path):
        # Add image to document
        doc.add_picture(str(flowchart_path), width=Inches(7.0))

        # Add caption
        caption = doc.add_paragraph()
        caption.add_run("Figure 1: ").bold = True
        caption.add_run("Data extraction pipeline flow chart. Gray cylinders represent database tables, "
                       "yellow boxes are intermediate RDS files, green boxes are final analysis datasets, "
                       "and blue rounded boxes are R extraction scripts.")
        caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].font.size = Pt(9)
        for run in caption.runs[1:]:
            run.font.size = Pt(9)
            run.font.italic = True
    else:
        doc.add_paragraph("(Flow chart could not be generated. Install Graphviz with: mamba install -n thesis graphviz)")

    # Save document
    doc.save(output_path)
    print(f"Created: {output_path}")


def main():
    parser = argparse.ArgumentParser(description='Generate logbook document for data extraction pipeline')
    parser.add_argument('--output', type=Path,
                        default=Path('suicidality/Documents/LOGBOOK_SSRI_Suicidality.docx'),
                        help='Output path for the logbook document')
    parser.add_argument('--project-name', type=str,
                        default='SSRI-Suicidality',
                        help='Project name for the logbook')

    args = parser.parse_args()

    # Ensure output directory exists
    args.output.parent.mkdir(parents=True, exist_ok=True)

    create_logbook_document(args.output, args.project_name)

    print(f"\nLogbook generated: {args.output}")


if __name__ == '__main__':
    main()
